import os

import cv2
import re
import pytesseract
from docx import Document
from utils.table_process.support_function import get_table_text,determine_table_position,\
    add_table_text,create_table_docx,create_cell_info
from utils.table_process.find_bb import IOU
from utils.main import entry_dla


def check_iou_with_above_row(cell1, cell2):
    (x, y, w, h) = cell1
    (x1, y1, w1, h1) = cell2
    if (y1 <= y <= y1 + h1 - 5) or (y <= y1 <= y + h - 5):
        return True
    return False


def check_iou_with_above_cell(cell1, cell2):
    (x, y, w, h) = cell1
    (x1, y1, w1, h1) = cell2
    if (x <= x1 <= x + w) or (x1 <= x <= x1 + w1):
        return True
    return False


def get_string_from_image(box, image, rule_base, auto_correct=True):
    (x, y, w, h) = box
    crop = image[y:y + h, x:x + w]
    string = pytesseract.image_to_string(crop, lang='vie')
    if auto_correct:
        string = rule_base.correct(string)
    cv2.rectangle(image, (x, y), (x + w, y + h), 255, 1)
    return string


def get_text_layout(list_result, list_big_box, img, doc_name, rule_base,
                    auto_correct=True):
    ## open or create docx file
    if not os.path.exists(doc_name):
        dc = Document()
        dc.save(doc_name)
    document = Document(doc_name)
    result = ""
    (height, _) = img.shape[:2]
	# if not have table
    if len(list_big_box) == 0:
        string = entry_dla(img, document, rule_base, auto_correct)
        document.save(doc_name)
        # string = pytesseract.image_to_string(img, lang='vie')
        # if auto_correct:
        # 	string = rule_base.correct(string)
        result.append(string + "\n")
        return result
	## if have table
    big_box_temp = []
    list_y_coord = [0]
    for (_, y, _, h) in list_big_box:
        big_box_temp.append((y, y + h))
        list_y_coord.append(y)
        list_y_coord.append(y + h)
    list_y_coord.append(height)
    for index, y1 in enumerate(list_y_coord):
        if index == len(list_y_coord) - 1:
            break
        y2 = list_y_coord[index + 1]
        if (y1, y2) not in big_box_temp:
            crop = img[y1:y2, :]  # not table
            idx = (crop.flatten() < 5)
            temp = sum(idx[:])
            if temp < 2:
                continue
            string = entry_dla(crop, document, rule_base, auto_correct)
            result = result + string # for get text
        else:
            index = 0
            box = list_big_box.pop(0)
            rows = []
            for temp in list_result:
                if IOU(temp[0], box):
                    index = index + 1
                    rows.append(temp)
                else:
                    break
            row_coords, col_coords, thresh = determine_table_position(rows)
            cell_infos = create_cell_info(rows, row_coords, col_coords, thresh)
            row = len(row_coords) - 1
            col = len(col_coords) - 1
            print("row", len(row_coords) - 1)
            print("col", len(col_coords) - 1)
            table = create_table_docx(cell_infos, row, col, document)
            string = add_table_text(table, rows, img, cell_infos)
            result = result + string
            list_result = list_result[index:]
    document.save(doc_name)
    result = normalize_text(result)
    return result

def normalize_text(text):
    text = re.sub('  ', '', text).strip()
    text = re.sub('"', '', text).strip()
    text = re.sub('\'', '', text).strip()
    text = re.sub(r'\\', '', text).strip()
    text = re.sub(r'\n', ' ', text).strip()
    text = re.sub(r'\r', ' ', text).strip()
    text = re.sub(r'\t', ' ', text).strip()
    text = re.sub('$', '', text).strip()
    return text

def get_text(list_result, list_big_box, img, rule_base, auto_correct=True):
    result = ""
    (height, _) = img.shape[:2]
    if len(list_big_box) == 0:
        result = result + pytesseract.image_to_string(img, lang='vie')
        return result
    big_box_temp = []
    list_y_coord = [0]
    for (_, y, _, h) in list_big_box:
        big_box_temp.append((y, y + h))
        list_y_coord.append(y)
        list_y_coord.append(y + h)
    list_y_coord.append(height)
    for index, y1 in enumerate(list_y_coord):
        if index == len(list_y_coord) - 1:
            break
        y2 = list_y_coord[index + 1]
        if (y1, y2) not in big_box_temp:
            crop = img[y1:y2, :]
            idx = (crop.flatten() < 5)
            temp = sum(idx[:])
            if temp < 2:
                continue
            result = result + pytesseract.image_to_string(crop, lang='vie')
        else:
            index = 0
            box = list_big_box.pop(0)
            rows = []
            for temp in list_result:
                if IOU(temp[0], box):
                    index = index + 1
                    rows.append(temp)
                else:
                    break
            row_coords, col_coords, thresh = determine_table_position(rows)
            cell_infos = create_cell_info(rows, row_coords, col_coords, thresh)
            print("row", len(row_coords) - 1)
            print("col", len(col_coords) - 1)
            string = get_table_text(rows, img, cell_infos)
            result= result + string
            list_result = list_result[index:]
    result = normalize_text(result)
    return result