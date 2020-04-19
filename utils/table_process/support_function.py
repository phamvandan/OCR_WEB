import cv2
import imutils
import pytesseract
from docx.enum.table import WD_TABLE_ALIGNMENT

def print_image(image):
    image = imutils.resize(image, width=500)
    cv2.imshow("ok",image)
    cv2.waitKey(0)

def is_append_coord(coords,coord,thresh):
    for val in coords:
        if abs(val-coord)<= thresh:
            return False
    return True


def determine_table_position(rows,delta=20):
    thresh = 0
    count = 0
    for row in rows:
        for index,cell in enumerate(row):
            if index==0:
                continue
            (x,y,w,h) = row[index-1]
            thresh = thresh + cell[0]-(x+w)
            count = count + 1
    if count > 0:
        thresh = thresh//count + delta
    else:
        thresh = delta
    print("table thresh",thresh)
    row_coords = []
    col_coords = []
    for row in rows:
        for cell in row:
            (x,y,w,h) = cell
            if is_append_coord(col_coords,x,thresh):
                col_coords.append(x)
            if is_append_coord(col_coords,x+w,thresh):
                col_coords.append(x+w)
            if is_append_coord(row_coords,y,thresh):
                row_coords.append(y)
            if is_append_coord(row_coords,y+h,thresh):
                row_coords.append(y+h)
    row_coords.sort(key=lambda x:x)
    # print(row_coords)
    col_coords.sort(key=lambda x:x)
    # del col_coords[0]
    # print("rows",len(row_coords))
    # print("cols",len(col_coords))
    return row_coords,col_coords,thresh


def create_cell_info(rows,row_coords,col_coords,thresh):
    cell_infos = []
    for row in rows:
        cell_info = []
        for cell in row:
            (x,y,w,h) = cell
            cell_row = []
            cell_col = []

            for index,col_coord in enumerate(col_coords):
                if len(cell_col)==2:
                    break
                if abs(x-col_coord)<=thresh or abs(x+w-col_coord)<=thresh:
                    cell_col.append(index)
                    continue

            for index,row_coord in enumerate(row_coords):
                if len(cell_row)==2:
                    break
                if abs(y-row_coord)<=thresh or abs(y+h-row_coord)<=thresh:
                    cell_row.append(index)
                    continue

            cell_row.sort(key=lambda x:x)
            cell_col.sort(key=lambda x:x)
            # print("cell row",cell_col)
            # print("cell col",cell_row)
            if len(cell_row)==2 and len(cell_col)==2:
                cell_info.append((cell_row[0],cell_col[0],cell_row[1]-1,cell_col[1]-1))
        cell_infos.append(cell_info)
    return cell_infos

def create_table_docx(cell_infos,row,col,document):
    ## create table
    table = document.add_table(rows=row, cols=col)
    ## merge cell
    for cell_info in cell_infos:
        for cell_position in cell_info:
            # print(cell_position)
            a = table.cell(cell_position[0],cell_position[1])
            b = table.cell(cell_position[2], cell_position[3])
            a.merge(b)
    return table

def add_table_text(table,rows,img,cell_infos):
    all_text = ""
    for index1,row in enumerate(rows):
        for index2,cell in enumerate(row):
            (x,y,w,h) = cell
            crop = img[y + 1:y + h - 1, x + 1:x + w - 1]
            string = ""
            if len(cell_infos[index1][index2])<2:
                continue
            (row,col) = cell_infos[index1][index2][:2]
            if col == 0:
                if row==0:
                    string = pytesseract.image_to_string(crop,config='--psm 8')
                else:
                    string = pytesseract.image_to_string(crop,config='--oem 3 --psm 6 outputbase digits')
            else:
                string = pytesseract.image_to_string(crop, lang='vie')
            row_cells = table.rows[row].cells
            p = row_cells[col].add_paragraph(string)
            p.alignment = WD_TABLE_ALIGNMENT.CENTER
            all_text = all_text + string + " "
    return all_text

def get_table_text(rows,img,cell_infos):
    all_text = ""
    for index1, row in enumerate(rows):
        for index2, cell in enumerate(row):
            (x, y, w, h) = cell
            crop = img[y+1:y + h-1, x+1:x + w-1]
            string = ""
            (row, col) = cell_infos[index1][index2][:2]
            if col == 0:
                if row == 0:
                    string = pytesseract.image_to_string(crop, config='--psm 8')
                else:
                    string = pytesseract.image_to_string(crop, config='--oem 3 --psm 6 outputbase digits')
            else:
                string = pytesseract.image_to_string(crop, lang='vie')
            all_text = all_text + string + " "
    return all_text

