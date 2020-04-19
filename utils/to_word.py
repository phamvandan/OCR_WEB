from utils.document_layout import layout_processing
import pytesseract
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


def get_string_from_image(box, image, rule_base, auto_correct=True,scale=2):
    (x, y, w, h) = box
    x = max(0,x-scale)
    y = max(0,y-scale)
    w = w + scale
    h = h + scale
    crop = image[y:y + h, x:x + w]
    text = pytesseract.image_to_string(crop, lang='vie',config='--psm 7')
    text = text.replace("\n", " ").strip()
    if auto_correct:
        text = rule_base.correct(text)
    return text
    # cv2.rectangle(image, (x, y), (x + w, y + h), 255, 1)


def layout_normal(line, image, table,rule_base, auto_correct,min_x):
    align = WD_TABLE_ALIGNMENT.LEFT
    row_cells = table.rows[0].cells
    region = line[0]
    string = ""
    for i,small_line in enumerate(region[1:]):
        ratio = (region[0][0] - min_x)/region[1][3]
        # print(ratio)
        if ratio>1:
            for i in range(0,int(ratio),2):
                string = string + "    "
        if small_line[0]-region[0][0]>small_line[3]:
            string = string + "       "
        string = string + get_string_from_image(small_line, image,rule_base, auto_correct) + "\n"
    string = string[:len(string) - 1]
    p = row_cells[0].add_paragraph(string)
    p.alignment = align
    return string

def layout_special(line,image,table,rule_base,auto_correct):
    align = WD_TABLE_ALIGNMENT.LEFT
    row_cells = table.rows[0].cells
    region = line[0]
    string = ""
    for i, small_line in enumerate(region[1:]):
        if small_line[0] - region[0][0] > small_line[3]:
            string = string + "       "
        string = string + get_string_from_image(small_line, image,rule_base, auto_correct) + "\n"
    string = string[:len(string) - 1]
    p = row_cells[1].add_paragraph(string)
    p.alignment = align
    return string


def find_min_x(lines):
    min_x = 10000
    max_x = 0
    for line in lines:
        if line[0][0][0]<min_x:
            min_x = line[0][0][0]
        if (line[0][0][0]+line[0][0][2])>max_x:
            max_x = line[0][0][0]+line[0][0][2]
    return min_x,max_x

def multiple_cols_table(line,image,table,rule_base, auto_correct=True):
    align = WD_TABLE_ALIGNMENT.CENTER
    row_cells = table.rows[0].cells
    region = line[0]
    all_text = ""
    for i, region in enumerate(line):
        string = ""
        for small_line in region[1:]:
            string = string + get_string_from_image(small_line, image,rule_base, auto_correct,scale=-1) + "\n"
        string = string[:len(string) - 1]
        p = row_cells[i].add_paragraph(string)
        p.alignment = align
        all_text = all_text + " " + string
    return all_text


def to_word(boxes, image, document, rule_base, auto_correct=True):
    lines = layout_processing(boxes, image)
    min_x, max_x = find_min_x(lines)
    print("minx", min_x)
    print("maxx", max_x)
    accept_distance = (max_x - min_x) // 20
    print("accept_distance ", accept_distance)
    all_text = ""
    for index, line in enumerate(lines):
        align = WD_TABLE_ALIGNMENT.CENTER
        column = len(line)
        string = ""
        if column>1:
            table = document.add_table(rows=1, cols=column)
            multiple_cols_table(line,image,table,rule_base,auto_correct)
        elif column == 1:
            (x, y, w, h) = line[0][1]
            ## case 2: if position in right page
            if (line[0][0][0])>=image.shape[1]//2:
                table = document.add_table(rows=1, cols=2)
                string = layout_special(line, image, table, rule_base, auto_correct)
                continue
            elif abs((max_x-x-w)-(x-min_x)) <= accept_distance:
                table = document.add_table(rows=1, cols=column)
                string = multiple_cols_table(line, image, table,rule_base,auto_correct)
            else:
                table = document.add_table(rows=1, cols=column)
                string = layout_normal(line, image, table,rule_base, auto_correct,min_x)
        # if column == 1:
        #     if (line[0][0][0])>=image.shape[1]//2:
        #         table = document.add_table(rows=1, cols=2)
        #         string = layout_special(line,image,table,rule_base,auto_correct)
        #         all_text = all_text + string
        #         continue
        # table = document.add_table(rows=1, cols=column)
        # if column == 1:
        #     if line[0][0][0] - min_x < 8*line[0][1][3]:
        #         string = layout_normal(line, image, table,rule_base, auto_correct,min_x)
        #         all_text = all_text + string
        #         continue
        # row_cells = table.rows[0].cells
        # for i, region in enumerate(line):
        #     string = ""
        #     for small_line in region[1:]:
        #         string = string + get_string_from_image(small_line, image, rule_base, auto_correct) + "\n"
        #     string = string[:len(string) - 1]
        #     all_text = all_text + string
        #     p = row_cells[i].add_paragraph(string)
        #     p.alignment = align
        all_text = all_text + string
    return all_text
